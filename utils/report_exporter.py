"""
报告导出工具
支持导出为HTML、Word、PDF等格式
"""
import logging
from typing import Dict, Any, List
from pathlib import Path
import json
from datetime import datetime
import tempfile
import base64

logger = logging.getLogger(__name__)

class ReportExporter:
    """报告导出器"""
    
    def __init__(self):
        self.supported_formats = ["html", "docx", "json"]
        # PDF导出需要额外的依赖，暂时不包含
    
    def export_report(self, report_data: Dict[str, Any], format_type: str, 
                     output_dir: str = None) -> Dict[str, Any]:
        """
        导出报告
        
        Args:
            report_data: 报告数据
            format_type: 导出格式 ('html', 'docx', 'json')
            output_dir: 输出目录
            
        Returns:
            Dict containing export result and file path
        """
        try:
            if format_type not in self.supported_formats:
                return {
                    "success": False,
                    "error": f"不支持的导出格式: {format_type}"
                }
            
            # 设置输出目录
            if not output_dir:
                output_dir = Path("outputs/reports")
            else:
                output_dir = Path(output_dir)
            
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_filename = f"conversation_report_{timestamp}"
            
            # 根据格式调用相应的导出方法
            if format_type == "html":
                result = self._export_html(report_data, output_dir, base_filename)
            elif format_type == "docx":
                result = self._export_docx(report_data, output_dir, base_filename)
            elif format_type == "json":
                result = self._export_json(report_data, output_dir, base_filename)
            
            logger.info(f"报告导出成功: {result.get('file_path')}")
            return result
            
        except Exception as e:
            logger.error(f"报告导出失败: {e}", exc_info=True)
            return {
                "success": False,
                "error": f"导出失败: {str(e)}"
            }
    
    def _export_html(self, report_data: Dict[str, Any], output_dir: Path, 
                    base_filename: str) -> Dict[str, Any]:
        """导出为HTML格式"""
        try:
            report = report_data.get("report", {})
            metadata = report.get("metadata", {})
            content = report.get("content", {})
            statistics = report.get("statistics", {})
            
            # 构建HTML内容
            html_content = self._build_html_content(metadata, content, statistics)
            
            # 保存文件
            file_path = output_dir / f"{base_filename}.html"
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(html_content)
            
            return {
                "success": True,
                "file_path": str(file_path),
                "format": "html",
                "file_size": file_path.stat().st_size
            }
            
        except Exception as e:
            raise Exception(f"HTML导出失败: {str(e)}")
    
    def _export_docx(self, report_data: Dict[str, Any], output_dir: Path, 
                    base_filename: str) -> Dict[str, Any]:
        """导出为Word文档格式"""
        try:
            # 尝试导入python-docx
            try:
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                return {
                    "success": False,
                    "error": "缺少python-docx依赖，请安装: pip install python-docx"
                }
            
            report = report_data.get("report", {})
            metadata = report.get("metadata", {})
            content = report.get("content", {})
            statistics = report.get("statistics", {})
            
            # 创建Word文档
            doc = Document()
            
            # 添加标题
            title = doc.add_heading(content.get("title", "对话报告"), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加元数据
            doc.add_heading("报告信息", level=1)
            info_table = doc.add_table(rows=0, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ("生成时间", metadata.get("generated_at", "")),
                ("分析类型", metadata.get("analysis_type", "")),
                ("对话轮数", str(metadata.get("conversation_count", 0))),
                ("报告样式", metadata.get("report_style", "")),
                ("生成器", metadata.get("generator", ""))
            ]
            
            for key, value in info_data:
                row_cells = info_table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = value
            
            # 添加执行摘要
            if content.get("executive_summary"):
                doc.add_heading("执行摘要", level=1)
                doc.add_paragraph(content["executive_summary"])
            
            # 添加主要内容
            if content.get("main_content"):
                doc.add_heading("详细内容", level=1)
                for section in content["main_content"]:
                    doc.add_heading(section.get("title", ""), level=2)
                    
                    if section.get("type") == "qa_pairs":
                        for i, qa in enumerate(section.get("content", []), 1):
                            doc.add_heading(f"问答 {i}", level=3)
                            doc.add_paragraph(f"问题：{qa.get('question', '')}")
                            doc.add_paragraph(f"回答：{qa.get('answer', '')}")
                            doc.add_paragraph("")  # 空行
                    
                    elif section.get("type") in ["bullet_list", "numbered_list"]:
                        for item in section.get("content", []):
                            p = doc.add_paragraph(item)
                            p.style = 'List Bullet' if section.get("type") == "bullet_list" else 'List Number'
                    
                    elif section.get("type") == "topics":
                        for topic in section.get("content", []):
                            doc.add_paragraph(f"• {topic}", style='List Bullet')
                    
                    elif section.get("type") == "keywords":
                        keywords_text = "、".join(section.get("content", []))
                        doc.add_paragraph(keywords_text)
            
            # 添加结论
            if content.get("conclusions"):
                conclusions = content["conclusions"]
                doc.add_heading("结论与建议", level=1)
                
                if conclusions.get("summary"):
                    doc.add_heading("总结", level=2)
                    doc.add_paragraph(conclusions["summary"])
                
                if conclusions.get("key_insights"):
                    doc.add_heading("关键洞察", level=2)
                    for insight in conclusions["key_insights"]:
                        doc.add_paragraph(f"• {insight}", style='List Bullet')
                
                if conclusions.get("recommendations"):
                    doc.add_heading("建议", level=2)
                    for recommendation in conclusions["recommendations"]:
                        doc.add_paragraph(f"• {recommendation}", style='List Bullet')
            
            # 添加统计信息
            if statistics:
                doc.add_heading("统计信息", level=1)
                stats_table = doc.add_table(rows=0, cols=2)
                stats_table.style = 'Table Grid'
                
                for key, value in statistics.items():
                    if isinstance(value, dict):
                        continue  # 跳过复杂对象
                    row_cells = stats_table.add_row().cells
                    row_cells[0].text = str(key)
                    row_cells[1].text = str(value)
            
            # 保存文件
            file_path = output_dir / f"{base_filename}.docx"
            doc.save(str(file_path))
            
            return {
                "success": True,
                "file_path": str(file_path),
                "format": "docx",
                "file_size": file_path.stat().st_size
            }
            
        except Exception as e:
            raise Exception(f"Word文档导出失败: {str(e)}")
    
    def _export_json(self, report_data: Dict[str, Any], output_dir: Path, 
                    base_filename: str) -> Dict[str, Any]:
        """导出为JSON格式"""
        try:
            # 保存文件
            file_path = output_dir / f"{base_filename}.json"
            with open(file_path, "w", encoding="utf-8") as f:
                json.dump(report_data, f, ensure_ascii=False, indent=2)
            
            return {
                "success": True,
                "file_path": str(file_path),
                "format": "json",
                "file_size": file_path.stat().st_size
            }
            
        except Exception as e:
            raise Exception(f"JSON导出失败: {str(e)}")
    
    def _build_html_content(self, metadata: Dict, content: Dict, statistics: Dict) -> str:
        """构建HTML内容"""
        html = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{content.get('title', '对话报告')}</title>
    <style>
        body {{
            font-family: 'Microsoft YaHei', Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        .container {{
            background-color: white;
            padding: 30px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            text-align: center;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            border-left: 4px solid #3498db;
            padding-left: 15px;
            margin-top: 30px;
        }}
        h3 {{
            color: #7f8c8d;
            margin-top: 25px;
        }}
        .metadata {{
            background-color: #ecf0f1;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
        }}
        .metadata table {{
            width: 100%;
            border-collapse: collapse;
        }}
        .metadata td {{
            padding: 8px;
            border-bottom: 1px solid #bdc3c7;
        }}
        .metadata td:first-child {{
            font-weight: bold;
            width: 30%;
        }}
        .qa-pair {{
            background-color: #f8f9fa;
            border: 1px solid #e9ecef;
            border-radius: 5px;
            padding: 15px;
            margin-bottom: 15px;
        }}
        .question {{
            color: #2980b9;
            font-weight: bold;
            margin-bottom: 10px;
        }}
        .answer {{
            color: #2c3e50;
            line-height: 1.7;
        }}
        .topics, .keywords {{
            display: flex;
            flex-wrap: wrap;
            gap: 8px;
        }}
        .tag {{
            background-color: #3498db;
            color: white;
            padding: 4px 12px;
            border-radius: 15px;
            font-size: 0.9em;
        }}
        .statistics-table {{
            width: 100%;
            border-collapse: collapse;
            margin-top: 15px;
        }}
        .statistics-table th, .statistics-table td {{
            border: 1px solid #bdc3c7;
            padding: 10px;
            text-align: left;
        }}
        .statistics-table th {{
            background-color: #34495e;
            color: white;
        }}
        ul {{
            padding-left: 20px;
        }}
        li {{
            margin-bottom: 5px;
        }}
        .print-button {{
            position: fixed;
            top: 20px;
            right: 20px;
            background-color: #3498db;
            color: white;
            border: none;
            padding: 10px 20px;
            border-radius: 5px;
            cursor: pointer;
            font-size: 14px;
        }}
        .print-button:hover {{
            background-color: #2980b9;
        }}
        @media print {{
            .print-button {{
                display: none;
            }}
            body {{
                background-color: white;
            }}
            .container {{
                box-shadow: none;
            }}
        }}
    </style>
</head>
<body>
    <button class="print-button" onclick="window.print()">打印报告</button>
    <div class="container">
        <h1>{content.get('title', '对话报告')}</h1>
        
        <div class="metadata">
            <h2>报告信息</h2>
            <table>
                <tr><td>生成时间</td><td>{metadata.get('generated_at', '')}</td></tr>
                <tr><td>分析类型</td><td>{metadata.get('analysis_type', '')}</td></tr>
                <tr><td>对话轮数</td><td>{metadata.get('conversation_count', 0)}</td></tr>
                <tr><td>报告样式</td><td>{metadata.get('report_style', '')}</td></tr>
                <tr><td>生成器</td><td>{metadata.get('generator', '')}</td></tr>
            </table>
        </div>
"""
        
        # 添加执行摘要
        if content.get("executive_summary"):
            html += f"""
        <h2>执行摘要</h2>
        <p>{content['executive_summary'].replace(chr(10), '<br>')}</p>
"""
        
        # 添加主要内容
        if content.get("main_content"):
            html += "<h2>详细内容</h2>"
            for section in content["main_content"]:
                html += f"<h3>{section.get('title', '')}</h3>"
                
                if section.get("type") == "qa_pairs":
                    for i, qa in enumerate(section.get("content", []), 1):
                        html += f"""
        <div class="qa-pair">
            <div class="question">问题 {i}：{qa.get('question', '')}</div>
            <div class="answer">{qa.get('answer', '').replace(chr(10), '<br>')}</div>
        </div>
"""
                
                elif section.get("type") in ["bullet_list", "numbered_list"]:
                    list_type = "ul" if section.get("type") == "bullet_list" else "ol"
                    html += f"<{list_type}>"
                    for item in section.get("content", []):
                        html += f"<li>{item}</li>"
                    html += f"</{list_type}>"
                
                elif section.get("type") == "topics":
                    html += '<div class="topics">'
                    for topic in section.get("content", []):
                        html += f'<span class="tag">{topic}</span>'
                    html += '</div>'
                
                elif section.get("type") == "keywords":
                    html += '<div class="keywords">'
                    for keyword in section.get("content", []):
                        html += f'<span class="tag">{keyword}</span>'
                    html += '</div>'
        
        # 添加结论
        if content.get("conclusions"):
            conclusions = content["conclusions"]
            html += "<h2>结论与建议</h2>"
            
            if conclusions.get("summary"):
                html += f"<h3>总结</h3><p>{conclusions['summary']}</p>"
            
            if conclusions.get("key_insights"):
                html += "<h3>关键洞察</h3><ul>"
                for insight in conclusions["key_insights"]:
                    html += f"<li>{insight}</li>"
                html += "</ul>"
            
            if conclusions.get("recommendations"):
                html += "<h3>建议</h3><ul>"
                for recommendation in conclusions["recommendations"]:
                    html += f"<li>{recommendation}</li>"
                html += "</ul>"
        
        # 添加统计信息
        if statistics:
            html += """
        <h2>统计信息</h2>
        <table class="statistics-table">
            <thead>
                <tr><th>指标</th><th>数值</th></tr>
            </thead>
            <tbody>
"""
            for key, value in statistics.items():
                if not isinstance(value, dict):
                    html += f"<tr><td>{key}</td><td>{value}</td></tr>"
            
            html += """
            </tbody>
        </table>
"""
        
        html += """
    </div>
</body>
</html>
"""
        return html
    
    def get_file_as_download(self, file_path: str) -> Dict[str, Any]:
        """获取文件的下载信息"""
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                return {
                    "success": False,
                    "error": "文件不存在"
                }
            
            # 读取文件内容
            with open(file_path, "rb") as f:
                file_content = f.read()
            
            # 编码为base64
            file_b64 = base64.b64encode(file_content).decode()
            
            return {
                "success": True,
                "filename": file_path.name,
                "content": file_b64,
                "size": len(file_content),
                "mime_type": self._get_mime_type(file_path.suffix)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"文件读取失败: {str(e)}"
            }
    
    def _get_mime_type(self, extension: str) -> str:
        """获取文件的MIME类型"""
        mime_types = {
            ".html": "text/html",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".json": "application/json",
            ".pdf": "application/pdf"
        }
        return mime_types.get(extension.lower(), "application/octet-stream")
